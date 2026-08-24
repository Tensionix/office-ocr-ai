# render_regions.py
# Render structured LayoutResult regions (from Surya / PP-Structure) into:
#   - DOCX: paragraphs/headings + REAL tables (explicit widths, merged spans)
#   - XLSX: real cell grid per table (merged spans)
# This is the reliable table path; use it instead of the flat-word heuristic
# whenever a layout provider is available.
# No globals; writes the given path. EN comments only. UTF-8 without BOM.
# Requires: python-docx, openpyxl.

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from openpyxl import Workbook
from openpyxl.styles import Font

from .engines.regions import LayoutResult, TableRegion, TextRegion


# --------------------------------------------------------------------------- #
# DOCX                                                                         #
# --------------------------------------------------------------------------- #

def regions_to_docx(out_path: str | Path,
                    pages: list[LayoutResult],
                    font_name: str = "DejaVu Sans",
                    page_width_in: float = 6.5) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)

    for pi, layout in enumerate(pages):
        if pi > 0:
            doc.add_page_break()
        for region in layout.regions:
            if isinstance(region, TextRegion):
                if not region.text.strip():
                    continue
                p = doc.add_paragraph(region.text)
                if region.kind == "heading":
                    p.style = doc.styles["Heading 1"]
            elif isinstance(region, TableRegion):
                _docx_table(doc, region, font_name, page_width_in)
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _docx_table(doc, region: TableRegion, font_name: str, page_width_in: float) -> None:
    if region.rows <= 0 or region.cols <= 0:
        return
    table = doc.add_table(rows=region.rows, cols=region.cols)
    table.style = "Table Grid"
    table.allow_autofit = False
    col_w = Inches(page_width_in / region.cols)  # equal split; explicit per-cell width

    grid = table.rows
    for c in region.cells:
        if c.row >= region.rows or c.col >= region.cols:
            continue
        cell = grid[c.row].cells[c.col]
        # merge spans
        if c.rowspan > 1 or c.colspan > 1:
            r2 = min(region.rows - 1, c.row + c.rowspan - 1)
            c2 = min(region.cols - 1, c.col + c.colspan - 1)
            cell = cell.merge(grid[r2].cells[c2])
        cell.text = c.text
        for col_cell in (cell.paragraphs[0].runs or [cell.paragraphs[0].add_run("")]):
            col_cell.font.name = font_name
    # explicit width on every cell (dual-width rule for reliable rendering)
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_w


# --------------------------------------------------------------------------- #
# XLSX                                                                         #
# --------------------------------------------------------------------------- #

def regions_to_xlsx(out_path: str | Path,
                    pages: list[LayoutResult],
                    font_name: str = "Arial") -> Path:
    wb = Workbook()
    wb.remove(wb.active)
    base_font = Font(name=font_name)

    for pi, layout in enumerate(pages):
        tregions = [r for r in layout.regions if isinstance(r, TableRegion)]
        if not tregions:
            continue
        ws = wb.create_sheet(title=f"page_{pi + 1}")
        row_cursor = 1
        for treg in tregions:
            row_cursor = _xlsx_table(ws, treg, base_font, row_cursor) + 2  # gap between tables

    if not wb.sheetnames:
        wb.create_sheet(title="page_1")
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out))
    return out


def _xlsx_table(ws, region: TableRegion, font: Font, start_row: int) -> int:
    last_row = start_row
    for c in region.cells:
        r = start_row + c.row
        col = c.col + 1
        cell = ws.cell(row=r, column=col, value=c.text)
        cell.font = font
        if c.rowspan > 1 or c.colspan > 1:
            ws.merge_cells(start_row=r, start_column=col,
                           end_row=r + c.rowspan - 1, end_column=col + c.colspan - 1)
        last_row = max(last_row, r + c.rowspan - 1)
    return last_row
