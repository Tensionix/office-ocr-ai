from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from docx_layout import (
    DocxLayoutOptions,
    add_docx_layout_arguments,
    apply_section_layout,
    describe_options,
    options_from_args,
)
from project_paths import ensure_project_dirs, iter_project_files, normalized_source_relative_path, output_path_for


INLINE_RE = re.compile(r"(`[^`]+`|\*\*.*?\*\*|\*.*?\*)")
BODY_FONT = "Arial"
BODY_SIZE_PT = 10.0
CODE_FONT = "Consolas"
DOCX_OPTIONS = DocxLayoutOptions(font_name=BODY_FONT, font_size_pt=BODY_SIZE_PT)


def apply_docx_options(options: DocxLayoutOptions) -> None:
    global BODY_FONT, BODY_SIZE_PT, DOCX_OPTIONS
    DOCX_OPTIONS = options
    BODY_FONT = options.font_name
    BODY_SIZE_PT = options.font_size_pt


def set_run_font_family(run, font_name: str) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def apply_run_style(
    run,
    *,
    font_name: str | None = None,
    size_pt: float | None = None,
    bold: bool = False,
    italic: bool = False,
) -> None:
    selected_font = font_name or BODY_FONT
    selected_size = BODY_SIZE_PT if size_pt is None else size_pt
    run.bold = bold
    run.italic = italic
    set_run_font_family(run, selected_font)
    run.font.size = Pt(selected_size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def parse_inline_formatting(paragraph, text: str) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            apply_run_style(run)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            apply_run_style(run, bold=True)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            apply_run_style(run, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            apply_run_style(run, font_name=CODE_FONT)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        apply_run_style(run)


def add_code_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    parse_inline_formatting(paragraph, text)
    for run in paragraph.runs:
        apply_run_style(run, font_name=CODE_FONT)


def configure_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    apply_section_layout(section, DOCX_OPTIONS)

    style_normal = doc.styles["Normal"]
    style_normal.font.name = BODY_FONT
    style_normal.font.size = Pt(BODY_SIZE_PT)
    style_normal.paragraph_format.space_after = Pt(0)

    for i in range(1, 4):
        style_heading = doc.styles[f"Heading {i}"]
        style_heading.font.name = BODY_FONT
        style_heading.font.color.rgb = RGBColor(0, 0, 0)
        style_heading.font.bold = True
        if i == 1:
            style_heading.font.size = Pt(14)
        elif i == 2:
            style_heading.font.size = Pt(12)
        else:
            style_heading.font.size = Pt(11)


def build_word_table(doc: Document, table_data: list[list[str]]) -> None:
    if not table_data:
        return
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row_data in enumerate(table_data):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= len(row_cells):
                continue
            paragraph = row_cells[col_idx].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            if row_idx == 0:
                clean_header = cell_text.replace("**", "").replace("*", "").replace("`", "")
                run = paragraph.add_run(clean_header)
                apply_run_style(run, bold=True)
            else:
                parse_inline_formatting(paragraph, cell_text)
            for run in paragraph.runs:
                apply_run_style(run)


def compile_markdown_to_docx_single(md_path: Path, out_path: Path, options: DocxLayoutOptions | None = None) -> None:
    ensure_project_dirs()
    if options is not None:
        apply_docx_options(options)
    doc = Document()
    configure_document_styles(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    table_buffer: list[list[str]] = []
    in_table = False
    in_code_block = False

    for raw_line in lines:
        line = raw_line.strip()
        if line.startswith("```"):
            if in_table:
                build_word_table(doc, table_buffer)
                table_buffer = []
                in_table = False
            in_code_block = not in_code_block
            continue

        if in_code_block:
            add_code_paragraph(doc, line)
            continue

        if line.startswith("|"):
            if "---" in line:
                continue
            parts = [part.strip() for part in line.split("|")[1:-1]]
            if parts:
                table_buffer.append(parts)
                in_table = True
            continue

        if in_table:
            build_word_table(doc, table_buffer)
            table_buffer = []
            in_table = False

        if not line:
            continue

        if line.startswith("# "):
            parse_inline_formatting(doc.add_heading(level=1), line[2:].strip())
        elif line.startswith("## "):
            parse_inline_formatting(doc.add_heading(level=2), line[3:].strip())
        elif line.startswith("### "):
            parse_inline_formatting(doc.add_heading(level=3), line[4:].strip())
        elif line.startswith("- ") or line.startswith("* "):
            parse_inline_formatting(doc.add_paragraph(style="List Bullet"), line[2:].strip())
        else:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            parse_inline_formatting(paragraph, line)

    if in_table:
        build_word_table(doc, table_buffer)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main_pipeline(options: DocxLayoutOptions | None = None) -> None:
    ensure_project_dirs()
    if options is not None:
        apply_docx_options(options)
    md_files = iter_project_files({".md"})
    if not md_files:
        print("[INFO] No .md files found in project root.")
        return

    print("=== AUDION OFFICE OCR AI: DOCX COMPILER (LLM) ===")
    print(f"Layout: {describe_options(DOCX_OPTIONS)}")
    for md_file in md_files:
        rel_path = normalized_source_relative_path(md_file)
        out_file = output_path_for(md_file, ".llm.docx")
        print(f"[BUILD] {rel_path} -> DOCX (LLM)")
        compile_markdown_to_docx_single(md_file, out_file)
    print("\n[OK] Compilation complete.")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Compile Markdown files to dense LLM DOCX.")
    add_docx_layout_arguments(parser, default_font_name=BODY_FONT, default_font_size_pt=BODY_SIZE_PT)
    return parser.parse_args()


if __name__ == "__main__":
    os.system("cls" if os.name == "nt" else "clear")
    main_pipeline(options_from_args(parse_args(), default_font_size_pt=BODY_SIZE_PT))
