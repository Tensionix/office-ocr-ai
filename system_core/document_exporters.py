"""Local exporters for the provider-neutral OCR DocumentModel package."""

from __future__ import annotations

import base64
import html
import json
import math
import re
import shutil
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

import fitz
from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from system_core.document_model import DocumentModel, write_document_model
from system_core.ocr_brick.engines.regions import TableRegion, TextRegion, region_from_dict
from system_core.ocr_brick.layout import build_grid, group_lines, line_text


EXPORT_SUFFIXES = {
    "docx": ".docx",
    "searchable_pdf": ".searchable.pdf",
    "xlsx": ".xlsx",
    "odt": ".odt",
    "markdown": ".md",
    "html": ".html",
    "json": ".document.json",
    "verification": ".verification.json",
    "archive": ".full-archive.zip",
}


def export_document_model(
    model: DocumentModel,
    package_dir: str | Path,
    output_dir: str | Path,
    formats: set[str],
    *,
    stem: str | None = None,
) -> dict[str, Path]:
    """Build any selected presentation formats locally, without OCR calls."""
    package = Path(package_dir)
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)
    base = stem or Path(model.source.name).stem
    produced: dict[str, Path] = {}
    ordinary = set(formats) - {"archive"}
    for fmt in sorted(ordinary):
        suffix = EXPORT_SUFFIXES.get(fmt)
        if not suffix:
            continue
        target = output / f"{base}{suffix}"
        if fmt == "docx":
            export_docx(model, package, target)
        elif fmt == "searchable_pdf":
            export_searchable_pdf(model, package, target)
        elif fmt == "xlsx":
            export_xlsx(model, target)
        elif fmt == "odt":
            export_odt(model, target)
        elif fmt == "markdown":
            export_markdown(model, target)
        elif fmt == "html":
            export_html(model, package, target)
        elif fmt == "json":
            shutil.copy2(package / "document.json", target)
        elif fmt == "verification":
            export_verification(model, target)
        produced[fmt] = target

    archive_target = output / f"{base}{EXPORT_SUFFIXES['archive']}" if "archive" in formats else None
    history_files = {key: path.name for key, path in produced.items()}
    if archive_target is not None:
        history_files["archive"] = archive_target.name
    model.export_history.append({
        "created_at": datetime.now(timezone.utc).isoformat(),
        "formats": sorted(history_files),
        "files": history_files,
    })
    write_document_model(model, package / "document.json")
    if "json" in produced:
        shutil.copy2(package / "document.json", produced["json"])
    if archive_target is not None:
        export_archive(package, archive_target, produced.values())
        produced["archive"] = archive_target
    return produced


def _regions(page) -> list[Any]:
    source = page.fused_regions if page.fused_regions else page.regions
    return [region_from_dict(item) for item in source if isinstance(item, dict)]


def _words(page) -> list[dict[str, Any]]:
    if page.fusion.get("applied"):
        for candidate in page.ocr_candidates:
            if candidate.get("engine") == "yandex" and isinstance(candidate.get("words"), list):
                return candidate["words"]
    if page.primary_words:
        return page.primary_words
    for candidate in page.ocr_candidates:
        words = candidate.get("words")
        if isinstance(words, list) and words:
            return words
    return []


def _paragraphs(page) -> list[str]:
    regions = _regions(page)
    text_regions = [item.text.strip() for item in regions if isinstance(item, TextRegion) and item.text.strip()]
    if text_regions:
        return text_regions
    words = _words(page)
    if words:
        return [line_text(line) for line in group_lines(words) if line_text(line).strip()]
    fallback = page.fused_text or page.primary_text
    return [part.strip() for part in fallback.splitlines() if part.strip()]


def _table_matrix(table: TableRegion) -> list[list[str]]:
    matrix = [["" for _ in range(max(1, table.cols))] for _ in range(max(1, table.rows))]
    for cell in table.cells:
        if 0 <= cell.row < len(matrix) and 0 <= cell.col < len(matrix[0]):
            matrix[cell.row][cell.col] = cell.text
    return matrix


def export_markdown(model: DocumentModel, target: str | Path) -> Path:
    lines = [f"# {Path(model.source.name).stem}", "", f"> OCR engine: `{model.metadata.get('engine', '')}`", ""]
    for page in model.pages:
        lines += [f"## Страница {page.index + 1}", ""]
        regions = _regions(page)
        if regions:
            for region in regions:
                if isinstance(region, TextRegion) and region.text.strip():
                    prefix = "### " if region.kind == "heading" else ""
                    lines += [prefix + region.text.strip(), ""]
                elif isinstance(region, TableRegion):
                    matrix = _table_matrix(region)
                    if matrix:
                        lines.append("| " + " | ".join(_md_cell(v) for v in matrix[0]) + " |")
                        lines.append("| " + " | ".join("---" for _ in matrix[0]) + " |")
                        lines += ["| " + " | ".join(_md_cell(v) for v in row) + " |" for row in matrix[1:]]
                        lines.append("")
        else:
            lines += [page.primary_text.strip(), ""]
    path = Path(target)
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return path


def _md_cell(value: str) -> str:
    return str(value).replace("|", "\\|").replace("\r", " ").replace("\n", "<br>")


def export_html(model: DocumentModel, package: Path, target: str | Path) -> Path:
    pages: list[str] = []
    for page in model.pages:
        image_path = package / page.source_page_image
        media = "image/png"
        image_uri = ""
        if image_path.is_file():
            image_uri = f"data:{media};base64,{base64.b64encode(image_path.read_bytes()).decode('ascii')}"
        body: list[str] = []
        for region in _regions(page):
            if isinstance(region, TextRegion) and region.text.strip():
                tag = "h3" if region.kind == "heading" else "p"
                body.append(f"<{tag}>{html.escape(region.text)}</{tag}>")
            elif isinstance(region, TableRegion):
                rows = "".join("<tr>" + "".join(f"<td>{html.escape(v)}</td>" for v in row) + "</tr>" for row in _table_matrix(region))
                body.append(f"<table>{rows}</table>")
        if not body:
            body.append(f"<pre>{html.escape(page.primary_text)}</pre>")
        pages.append(f'<section><h2>Страница {page.index + 1}</h2><div class="page-grid"><img src="{image_uri}" alt="Страница {page.index + 1}"><article>{"".join(body)}</article></div></section>')
    doc = f"""<!doctype html><html lang="ru"><head><meta charset="utf-8"><title>{html.escape(model.source.name)}</title>
<style>body{{font:16px Arial,sans-serif;margin:24px;color:#202124}}section{{margin:0 0 40px}}.page-grid{{display:grid;grid-template-columns:minmax(320px,1fr) minmax(320px,1fr);gap:24px;align-items:start}}img{{max-width:100%;box-shadow:0 2px 12px #999}}table{{border-collapse:collapse;width:100%}}td{{border:1px solid #777;padding:5px;vertical-align:top}}pre{{white-space:pre-wrap}}@media(max-width:900px){{.page-grid{{grid-template-columns:1fr}}}}</style></head><body><h1>{html.escape(Path(model.source.name).stem)}</h1>{''.join(pages)}</body></html>"""
    path = Path(target)
    path.write_text(doc, encoding="utf-8")
    return path


def export_docx(model: DocumentModel, package: Path, target: str | Path) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    for index, page in enumerate(model.pages):
        if index:
            doc.add_page_break()
        section = doc.sections[-1]
        page_regions = _regions(page)
        max_cols = max((r.cols for r in page_regions if isinstance(r, TableRegion)), default=0)
        landscape = page.width > page.height or max_cols > 7
        if landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        if max_cols > 12:
            section.page_width, section.page_height = Inches(16.5), Inches(11.7)  # A3 landscape
        margin = 0.35 if max_cols > 12 else 0.55
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)
        usable_width = float(section.page_width.inches - section.left_margin.inches - section.right_margin.inches)
        regions = page_regions
        if not regions:
            for text in _paragraphs(page):
                doc.add_paragraph(text)
            continue
        for region in regions:
            if isinstance(region, TextRegion) and region.text.strip():
                para = doc.add_paragraph(region.text)
                if region.kind == "heading":
                    para.style = doc.styles["Heading 2"]
            elif isinstance(region, TableRegion):
                _docx_table(doc, region, usable_width)
            elif getattr(region, "kind", "") == "figure" and region.box:
                _docx_figure(doc, package / page.source_page_image, region.box)
    path = Path(target)
    doc.save(path)
    return path


def _docx_table(doc: Document, region: TableRegion, usable_width: float) -> None:
    if region.rows < 1 or region.cols < 1:
        return
    table = doc.add_table(rows=region.rows, cols=region.cols)
    table.style = "Table Grid"
    table.autofit = False
    widths = _balanced_column_widths(region, usable_width)
    for col_index, column in enumerate(table.columns):
        column.width = Inches(widths[col_index])
        for cell in column.cells:
            cell.width = Inches(widths[col_index])
    for item in region.cells:
        if not (0 <= item.row < region.rows and 0 <= item.col < region.cols):
            continue
        cell = table.cell(item.row, item.col)
        if item.rowspan > 1 or item.colspan > 1:
            cell = cell.merge(table.cell(min(region.rows - 1, item.row + item.rowspan - 1), min(region.cols - 1, item.col + item.colspan - 1)))
        cell.text = item.text
        spanned_width = sum(widths[item.col:min(region.cols, item.col + max(1, item.colspan))])
        cell.width = Inches(spanned_width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(6.5 if region.cols > 12 else (7.5 if region.cols > 7 else 9))
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.first_child_found_in("w:tcW")
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
        tc_w.set(qn("w:w"), str(int(spanned_width * 1440)))
        tc_w.set(qn("w:type"), "dxa")
        if tc_w.getparent() is None:
            tc_pr.append(tc_w)


_NUMERIC_CELL_RE = re.compile(r"^[\s\d.,%+\-–—/()№]+$")


def _balanced_column_widths(region: TableRegion, total_width: float) -> list[float]:
    """Allocate real page width by content while keeping every column usable."""
    if region.cols <= 0:
        return []
    scores = [4.0] * region.cols
    for item in region.cells:
        if item.colspan != 1 or not (0 <= item.col < region.cols):
            continue
        text = " ".join(str(item.text or "").split())
        if not text:
            continue
        longest = max((len(part) for part in re.split(r"\s+", text)), default=1)
        if _NUMERIC_CELL_RE.fullmatch(text):
            score = 3.0 if len(text) <= 8 else 4.5
        elif len(text) <= 6 and longest <= 6:
            score = 4.0
        else:
            score = min(22.0, 4.5 + math.sqrt(min(len(text), 500)) * 0.8 + min(longest, 30) * 0.18)
        scores[item.col] = max(scores[item.col], score)

    minimum = 0.28 if region.cols > 20 else (0.38 if region.cols > 10 else 0.55)
    maximum = min(
        total_width - minimum * (region.cols - 1),
        max(total_width / region.cols * 1.6, min(3.5, max(1.25, total_width * 0.24))),
    )
    widths = [minimum] * region.cols
    remaining = max(0.0, total_width - sum(widths))
    active = set(range(region.cols))
    while remaining > 1e-8 and active:
        total_score = sum(scores[i] for i in active)
        capped: list[int] = []
        for i in active:
            share = remaining * scores[i] / total_score
            if share >= maximum - widths[i] - 1e-8:
                capped.append(i)
        if not capped:
            for i in active:
                widths[i] += remaining * scores[i] / total_score
            remaining = 0.0
            break
        for i in capped:
            addition = max(0.0, maximum - widths[i])
            widths[i] += addition
            remaining -= addition
            active.remove(i)
    if remaining > 1e-8:
        widths[-1] += remaining
    return widths


def _docx_figure(doc: Document, image_path: Path, box: tuple[int, int, int, int]) -> None:
    if not image_path.is_file():
        return
    crop_path = image_path.with_suffix(".figure.tmp.png")
    try:
        with Image.open(image_path) as image:
            x, y, w, h = box
            image.crop((x, y, x + w, y + h)).save(crop_path)
        doc.add_picture(str(crop_path), width=Inches(min(6.8, max(1.0, w / 300))))
    finally:
        crop_path.unlink(missing_ok=True)


def export_searchable_pdf(model: DocumentModel, package: Path, target: str | Path) -> Path:
    pdf = fitz.open()
    font_path = Path(r"C:\Windows\Fonts\arial.ttf")
    for page_model in model.pages:
        image_path = package / page_model.source_page_image
        if not image_path.is_file():
            continue
        with Image.open(image_path) as image:
            width_px, height_px = image.size
        dpi = max(1, page_model.dpi or 300)
        width_pt, height_pt = width_px * 72 / dpi, height_px * 72 / dpi
        page = pdf.new_page(width=width_pt, height=height_pt)
        page.insert_image(page.rect, filename=str(image_path))
        fontname = "ocrfont"
        if font_path.is_file():
            page.insert_font(fontname=fontname, fontfile=str(font_path))
        else:
            fontname = "helv"
        sx = width_pt / max(1, page_model.width or width_px)
        sy = height_pt / max(1, page_model.height or height_px)
        words = _words(page_model)
        for word in words:
            text = str(word.get("text") or "").strip()
            box = word.get("box")
            if not text or not isinstance(box, (list, tuple)) or len(box) != 4:
                continue
            x, y, w, h = [float(v) for v in box]
            page.insert_text(
                fitz.Point(x * sx, (y + h * 0.82) * sy),
                text,
                fontname=fontname,
                fontsize=max(2.5, h * sy * 0.78),
                render_mode=3,
                overlay=True,
            )
        if not words and page_model.primary_text.strip():
            # Provider text without word boxes remains searchable; coordinates are
            # approximate and the exact visual layer still comes from the source.
            page.insert_textbox(
                fitz.Rect(2, 2, width_pt - 2, height_pt - 2),
                page_model.primary_text,
                fontname=fontname,
                fontsize=6,
                render_mode=3,
                overlay=True,
            )
    path = Path(target)
    pdf.save(path, deflate=True, garbage=4)
    pdf.close()
    return path


def export_xlsx(model: DocumentModel, target: str | Path) -> Path:
    wb = Workbook()
    wb.remove(wb.active)
    thin = Side(style="thin", color="B7C3D0")
    for page in model.pages:
        tables = [r for r in _regions(page) if isinstance(r, TableRegion)]
        sheet = wb.create_sheet(f"Страница {page.index + 1}"[:31])
        row_cursor = 1
        if not tables:
            grid = build_grid(_words(page))
            tables = []
            for r, row in enumerate(grid, 1):
                for c, value in enumerate(row, 1):
                    sheet.cell(r, c, value)
            row_cursor = len(grid) + 1
        for table_index, table in enumerate(tables, 1):
            sheet.cell(row_cursor, 1, f"Таблица {table_index}").font = Font(name="Arial", bold=True, color="FFFFFF")
            sheet.cell(row_cursor, 1).fill = PatternFill("solid", fgColor="2F5597")
            start = row_cursor + 1
            for item in table.cells:
                r, c = start + item.row, item.col + 1
                cell = sheet.cell(r, c, item.text)
                cell.font = Font(name="Arial", size=10, bold=item.row == 0)
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
                if item.row == 0:
                    cell.fill = PatternFill("solid", fgColor="D9EAF7")
                if item.rowspan > 1 or item.colspan > 1:
                    sheet.merge_cells(start_row=r, start_column=c, end_row=r + item.rowspan - 1, end_column=c + item.colspan - 1)
            row_cursor = start + table.rows + 2
        max_col = max(1, sheet.max_column)
        for col in range(1, max_col + 1):
            values = [str(sheet.cell(row, col).value or "") for row in range(1, min(sheet.max_row, 200) + 1)]
            nonempty = [value for value in values if value]
            numeric = bool(nonempty) and sum(bool(_NUMERIC_CELL_RE.fullmatch(value)) for value in nonempty) / len(nonempty) >= 0.7
            longest_word = max((len(word) for value in nonempty for word in value.split()), default=0)
            if numeric:
                width = min(16, max(9, max((len(value) for value in nonempty), default=0) + 2))
            else:
                width = min(45, max(12, longest_word + 3, math.sqrt(max((len(value) for value in nonempty), default=0)) * 3.2))
            sheet.column_dimensions[get_column_letter(col)].width = width
        sheet.freeze_panes = "A2"
        sheet.sheet_view.showGridLines = False
    audit = wb.create_sheet("OCR Audit")
    audit.append([
        "Страница", "Основной провайдер", "Tesseract статус", "Tesseract числа",
        "Yandex статус", "Yandex числа", "Fusion замен", "Fusion проверить", "Confidence",
    ])
    for cell in audit[1]:
        cell.font = Font(name="Arial", bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="2F5597")
    for page in model.pages:
        verification = page.verification
        yandex_verification = page.provider_verifications.get("yandex", {})
        audit.append([
            page.index + 1,
            page.primary_engine,
            verification.get("status", ""),
            verification.get("numeric_agreement", ""),
            yandex_verification.get("status", ""),
            yandex_verification.get("numeric_agreement", ""),
            page.fusion.get("replaced_items", 0),
            page.fusion.get("review_items", 0),
            _mean_confidence(_words(page)),
        ])
    audit.freeze_panes = "A2"
    for col, width in enumerate((12, 22, 20, 18, 18, 16, 16, 18, 16), 1):
        audit.column_dimensions[get_column_letter(col)].width = width
    path = Path(target)
    wb.save(path)
    return path


def _mean_confidence(words: list[dict[str, Any]]) -> float | str:
    values = [float(w["confidence"]) for w in words if isinstance(w.get("confidence"), (int, float)) and float(w["confidence"]) >= 0]
    return round(sum(values) / len(values), 4) if values else ""


def export_odt(model: DocumentModel, target: str | Path) -> Path:
    body: list[str] = []
    for page in model.pages:
        body.append(f'<text:h text:outline-level="1">Страница {page.index + 1}</text:h>')
        regions = _regions(page)
        if not regions:
            body += [f"<text:p>{escape(text)}</text:p>" for text in _paragraphs(page)]
        for region in regions:
            if isinstance(region, TextRegion) and region.text.strip():
                body.append(f"<text:p>{escape(region.text)}</text:p>")
            elif isinstance(region, TableRegion):
                body.append('<table:table table:name="OCRTable">')
                for row in _table_matrix(region):
                    body.append("<table:table-row>" + "".join(f'<table:table-cell office:value-type="string"><text:p>{escape(value)}</text:p></table:table-cell>' for value in row) + "</table:table-row>")
                body.append("</table:table>")
    content = f'''<?xml version="1.0" encoding="UTF-8"?><office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0" office:version="1.3"><office:body><office:text>{''.join(body)}</office:text></office:body></office:document-content>'''
    styles = '''<?xml version="1.0" encoding="UTF-8"?><office:document-styles xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" office:version="1.3"><office:styles/></office:document-styles>'''
    manifest = '''<?xml version="1.0" encoding="UTF-8"?><manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.3"><manifest:file-entry manifest:full-path="/" manifest:media-type="application/vnd.oasis.opendocument.text"/><manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="styles.xml" manifest:media-type="text/xml"/></manifest:manifest>'''
    path = Path(target)
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("mimetype", "application/vnd.oasis.opendocument.text", compress_type=zipfile.ZIP_STORED)
        archive.writestr("content.xml", content)
        archive.writestr("styles.xml", styles)
        archive.writestr("META-INF/manifest.xml", manifest)
    return path


def export_verification(model: DocumentModel, target: str | Path) -> Path:
    payload = [{
        "page": page.index + 1,
        "primary_verification": page.verification,
        "provider_verifications": page.provider_verifications,
        "fusion": page.fusion,
    } for page in model.pages]
    path = Path(target)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return path


def export_archive(package: Path, target: str | Path, exports: Any = ()) -> Path:
    path = Path(target)
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as archive:
        for item in sorted(package.rglob("*")):
            if item.is_file():
                archive.write(item, Path("DocumentModel") / item.relative_to(package))
        for item in exports:
            item = Path(item)
            if item.is_file():
                archive.write(item, Path("Exports") / item.name)
    return path
