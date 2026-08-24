from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Mm, Pt, RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from docx_layout import (
    DocxLayoutOptions,
    add_docx_layout_arguments,
    apply_section_layout,
    describe_options,
    options_from_args,
)
from project_paths import ensure_project_dirs, iter_project_files, normalized_source_relative_path, output_path_for

BODY_FONT = "Arial"
BODY_SIZE_PT = 10.5
CODE_FONT = "Consolas"
CODE_SIZE_PT = 10.5
HEADING_1_PT = 20
HEADING_2_PT = 16
HEADING_3_PT = 15
DOCX_OPTIONS = DocxLayoutOptions(font_name=BODY_FONT, font_size_pt=BODY_SIZE_PT)

URL_RE = re.compile(r"https?://[^\s)]+")
INLINE_RE = re.compile(r"(`[^`]+`|\*\*.+?\*\*|\*.+?\*)")
ORDERED_RE = re.compile(r"^\d+\.\s+")
TAG_LINE_RE = re.compile(r"^\s*(#[^\s]+)(?:\s+#[^\s]+)*\s*$")


def apply_docx_options(options: DocxLayoutOptions) -> None:
    global BODY_FONT, BODY_SIZE_PT, DOCX_OPTIONS
    DOCX_OPTIONS = options
    BODY_FONT = options.font_name
    BODY_SIZE_PT = options.font_size_pt


def set_run_font(run, *, font_name: str, size_pt: float, bold: bool = False, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def set_run_shading(run, fill: str) -> None:
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    rpr.append(shd)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)

    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rfonts.set(qn("w:cs"), BODY_FONT)
    rpr.append(rfonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(BODY_SIZE_PT * 2)))
    rpr.append(sz)

    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), str(int(BODY_SIZE_PT * 2)))
    rpr.append(szcs)

    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def ensure_preview_style(doc: Document) -> None:
    if "code-line" in doc.styles:
        style = doc.styles["code-line"]
    else:
        style = doc.styles.add_style("code-line", WD_STYLE_TYPE.PARAGRAPH)

    style.font.name = BODY_FONT
    style.font.size = Pt(BODY_SIZE_PT)
    style.paragraph_format.space_before = Pt(5)
    style.paragraph_format.space_after = Pt(5)
    style.paragraph_format.line_spacing = 1.0


def configure_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    apply_section_layout(section, DOCX_OPTIONS)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = BODY_FONT
    heading1.font.size = Pt(HEADING_1_PT)
    heading1.font.bold = True
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(12)

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = BODY_FONT
    heading2.font.size = Pt(HEADING_2_PT)
    heading2.font.bold = True
    heading2.paragraph_format.space_before = Pt(18)
    heading2.paragraph_format.space_after = Pt(12)

    heading3 = doc.styles["Heading 3"]
    heading3.font.name = BODY_FONT
    heading3.font.size = Pt(HEADING_3_PT)
    heading3.font.bold = True
    heading3.paragraph_format.space_before = Pt(18)
    heading3.paragraph_format.space_after = Pt(12)

    ensure_preview_style(doc)


def apply_code_line_paragraph_format(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.first_line_indent = None


def parse_inline_runs(paragraph, text: str) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            add_text_with_urls(paragraph, text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            add_plain_text(paragraph, token[2:-2], bold=True)
        elif token.startswith("*") and token.endswith("*"):
            add_plain_text(paragraph, token[1:-1], italic=True)
        else:
            add_code_text(paragraph, token[1:-1])
        pos = match.end()
    if pos < len(text):
        add_text_with_urls(paragraph, text[pos:])


def add_plain_text(paragraph, text: str, *, bold: bool = False, italic: bool = False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    set_run_font(run, font_name=BODY_FONT, size_pt=BODY_SIZE_PT, bold=bold, italic=italic)


def add_code_text(paragraph, text: str) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    set_run_font(run, font_name=CODE_FONT, size_pt=CODE_SIZE_PT)
    set_run_shading(run, "D9D9D9")


def add_text_with_urls(paragraph, text: str) -> None:
    pos = 0
    for match in URL_RE.finditer(text):
        if match.start() > pos:
            add_plain_text(paragraph, text[pos:match.start()])
        url = match.group(0)
        add_hyperlink(paragraph, url, url)
        pos = match.end()
    if pos < len(text):
        add_plain_text(paragraph, text[pos:])


def add_tag_line(paragraph, text: str) -> None:
    parts = re.split(r"(\s+)", text.strip())
    for part in parts:
        if not part:
            continue
        if part.isspace():
            add_plain_text(paragraph, "\u00a0")
        elif part.startswith("#"):
            add_code_text(paragraph, part)
        else:
            add_plain_text(paragraph, part)


def build_word_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    norm_rows = [row + [""] * (col_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(norm_rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row in enumerate(norm_rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.0
            if row_idx == 0:
                parse_inline_runs(paragraph, value.replace("**", "").replace("*", ""))
                for run in paragraph.runs:
                    run.bold = True
            else:
                parse_inline_runs(paragraph, value)
    doc.add_paragraph()


def normalize_source_line(text: str) -> str:
    stripped = text.strip()
    if len(stripped) >= 2 and stripped.startswith("_") and stripped.endswith("_"):
        return stripped[1:-1].strip()
    return stripped


def merge_followup_link_lines(blocks: list[tuple[str, object]]) -> list[tuple[str, object]]:
    merged: list[tuple[str, object]] = []
    for kind, payload in blocks:
        if (
            kind == "paragraph"
            and isinstance(payload, list)
            and len(payload) == 1
            and str(payload[0]).startswith("Ссылка: ")
            and merged
        ):
            prev_kind, prev_payload = merged[-1]
            if prev_kind == "paragraph" and isinstance(prev_payload, list):
                prev_payload.append(str(payload[0]))
                continue
            if prev_kind in {"ordered_item", "bullet_item"} and isinstance(prev_payload, str):
                merged[-1] = ("paragraph", [prev_payload, str(payload[0])])
                continue
        merged.append((kind, payload))
    return merged


def preprocess_blocks(lines: list[str]) -> list[tuple[str, object]]:
    blocks: list[tuple[str, object]] = []
    paragraph_lines: list[str] = []
    table_rows: list[list[str]] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            blocks.append(("paragraph", paragraph_lines[:]))
            paragraph_lines = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            blocks.append(("table", table_rows[:]))
            table_rows = []

    for raw in lines:
        line = raw.rstrip()
        stripped = normalize_source_line(line)

        if stripped.startswith("|"):
            flush_paragraph()
            if "---" not in stripped:
                table_rows.append([part.strip() for part in stripped.split("|")[1:-1]])
            continue

        flush_table()

        if not stripped:
            flush_paragraph()
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            blocks.append(("heading1", stripped[2:].strip()))
            continue
        if stripped.startswith("## "):
            flush_paragraph()
            blocks.append(("heading2", stripped[3:].strip()))
            continue
        if stripped.startswith("### "):
            flush_paragraph()
            blocks.append(("heading3", stripped[4:].strip()))
            continue

        if ORDERED_RE.match(stripped):
            flush_paragraph()
            blocks.append(("ordered_item", ORDERED_RE.sub("", stripped, count=1)))
            continue

        if stripped.startswith(("- ", "* ")):
            flush_paragraph()
            blocks.append(("bullet_item", stripped[2:].strip()))
            continue

        paragraph_lines.append(stripped)

    flush_table()
    flush_paragraph()
    return merge_followup_link_lines(blocks)


def render_code_line_paragraph(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph(style="code-line")
    apply_code_line_paragraph_format(paragraph)

    for line_index, line in enumerate(lines):
        if line_index > 0:
            paragraph.add_run().add_break()

        if TAG_LINE_RE.match(line):
            add_tag_line(paragraph, line)
        else:
            parse_inline_runs(paragraph, line)


def compile_markdown_to_docx_single(md_path: Path, out_path: Path, options: DocxLayoutOptions | None = None) -> None:
    ensure_project_dirs()
    if options is not None:
        apply_docx_options(options)
    doc = Document()
    configure_document_styles(doc)

    blocks = preprocess_blocks(md_path.read_text(encoding="utf-8").splitlines())

    for kind, payload in blocks:
        if kind == "heading1":
            paragraph = doc.add_heading(str(payload), level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if kind == "heading2":
            paragraph = doc.add_heading(str(payload), level=2)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if kind == "heading3":
            paragraph = doc.add_heading(str(payload), level=3)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if kind == "table":
            build_word_table(doc, payload)  # type: ignore[arg-type]
            continue
        if kind in {"ordered_item", "bullet_item"}:
            render_code_line_paragraph(doc, [str(payload)])
            continue
        render_code_line_paragraph(doc, payload)  # type: ignore[arg-type]

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main_pipeline(options: DocxLayoutOptions | None = None) -> None:
    ensure_project_dirs()
    if options is not None:
        apply_docx_options(options)
    md_files = iter_project_files({".md"})
    if not md_files:
        print("[INFO] No .md files found in project root.")
        return

    print("=== AUDION OFFICE OCR AI: DOCX COMPILER ===")
    print(f"Layout: {describe_options(DOCX_OPTIONS)}")
    for md_file in md_files:
        rel_path = normalized_source_relative_path(md_file)
        out_file = output_path_for(md_file, ".docx")
        print(f"[BUILD] {rel_path} -> DOCX")
        compile_markdown_to_docx_single(md_file, out_file)
    print("\n[OK] Compilation complete.")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Compile Markdown files to DOCX.")
    add_docx_layout_arguments(parser, default_font_name=BODY_FONT, default_font_size_pt=BODY_SIZE_PT)
    return parser.parse_args()


if __name__ == "__main__":
    os.system("cls" if os.name == "nt" else "clear")
    main_pipeline(options_from_args(parse_args(), default_font_size_pt=BODY_SIZE_PT))
